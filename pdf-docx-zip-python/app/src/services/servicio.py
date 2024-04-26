from io import BytesIO
from docx import Document
import fitz
import os
import tempfile

EMPTY_TEXT = "" #Para cuando el archivo esté vacío
class Metodos:
    @staticmethod
    def text_from_pdf(file): 
        """Extrae de un archivo PDF (entero) el texto y lo devuelve en una cadena de texto plano."""
        text = EMPTY_TEXT
        with fitz.open(file) as doc:
            for page in doc:
                text += page.get_text()
        return text
    
    @staticmethod
    def text_from_temp_pdf(pdf_data): 
        """Extrae de los datos de un PDF en memoria (entero) a texto, crea un doc Fitz y lo devuelve en una cadena de texto plano"""
        text = EMPTY_TEXT
        temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) # A False para poder abrirlo con fitz primero
        temp_pdf.write(pdf_data.read())
        temp_pdf.close()
        with fitz.open(temp_pdf.name) as doc:
            for page in doc:
                text += page.get_text()  
        os.unlink(temp_pdf.name) #Borra el temp_pdf
        return text
    
    @staticmethod
    def text_from_pdf_page(page_num, file): 
        """Extrae de un archivo pdf el texto de una página específica y lo devuelve en una cadena de texto plano."""
        text = EMPTY_TEXT
        with fitz.open(file) as doc:
            page = doc.load_page(page_num)
            text = page.get_text()

        return text

    @staticmethod
    def text_from_pdf_pages(page_start, page_finish, file):
        """Extrae de un archivo pdf el texto de un rango de páginas específicas (ambas incluidas) y lo devuelve en una cadena de texto plano."""
        text = EMPTY_TEXT
        with fitz.open(file) as doc:
            total_pages = len(doc)
            if page_start < 1 or page_start > total_pages or page_finish < 1 or page_finish > total_pages:
                raise ValueError("Páginas fuera de rango")
            
            for page_num in range(page_start -1, page_finish):
                page = doc.load_page(page_num)
                text += page.get_text() + '\n'
        return text
    
    @staticmethod
    def transform_pdf_to_png(page_num, file):
        """Transforma de un archivo pdf una página específica a un pixmap en 300 dpi y lo devuelve."""
        with fitz.open(file) as doc:
            page = doc.load_page(page_num)
            pixmap = page.get_pixmap(dpi=300, alpha=False) #alpha = false, sin transparencia

        return pixmap
    
    @staticmethod
    def transform_entire_pdf_to_png(file):
        """Transforma de un archivo pdf a una lista de pixmap en 300 dpi y lo devuelve."""
        pixmap_list = []
        with fitz.open(file) as doc:
            for page in doc:
                pixmap = page.get_pixmap(dpi=300, alpha=False) #alpha = false, sin transparencia
                pixmap_list.append(pixmap)
        return pixmap_list
       
    @staticmethod
    def extract_img_from_pdf(file):
        """Extrae de archivo pdf todas las páginas y las devuelve en una lista de bytestream."""    
        imgs = []
        with fitz.open(file) as doc:
            for page_num in range(len(doc)): #Itera por todas las páginas
                try:
                    page = doc.load_page(page_num)
                    image_list = page.get_images(full=True) #full=True, se obtiene una lista de tuplas con información detallada
                    for img in enumerate(image_list):
                        for img_tupla in image_list:
                            print(img_tupla)
                            xref = img_tupla[0]
                            dict_image = doc.extract_image(xref)
                            image_stream = BytesIO(dict_image["image"])
                            imgs.append(image_stream)
                except Exception as e:
                    print(f"Error en la página {page_num + 1}: {e}")
        return imgs
    
    @staticmethod
    def extract_img_from_pdf_pages(page_start, page_finish, file):
        """Transforma de un archivo pdf un rango de páginas específicas a pixmap en 300 dpi y las devuelve en una lista de bytestream."""
        imgs = []
        with fitz.open(file) as doc:
            total_pages = len(doc)
            if page_start < 1 or page_start > total_pages or page_finish < 1 or page_finish > total_pages:
                raise ValueError("Páginas fuera de rango")     
            
            for page_num in range(page_start -1, page_finish): #Itera por la pagins proporcionadas
                try:
                    page = doc.load_page(page_num)
                    image_list = page.get_images(full=True) #full=True, se obtiene una lista de tuplas con información detallada

                    for img in enumerate(image_list):
                        for img_tupla in image_list:
                            print(img_tupla)
                            xref = img_tupla[0]
                            dict_image = doc.extract_image(xref)
                            image_stream = BytesIO(dict_image["image"])
                            imgs.append(image_stream)
                except Exception as e:
                    print(f"Error en la página {page_num + 1}: {e}")
        return imgs

    @staticmethod
    def extract_text_from_docx(file):
        """Extrae de un archivo docx el texto de todas sus páginas y lo devuelve en una cadena de texto plano."""
        document = Document(file)
        text = EMPTY_TEXT
        for parag in document.paragraphs:
            text += parag.text + "\n"
        return text
    
    @staticmethod
    def create_docx_from_text(text, out_dir, filename='documento.docx'):
        """Crea un docx a partir de un texto plano con el nombre porporcionado y lo guarda en la dirección proporcionada."""
        filename = os.path.splitext(filename)[0] + ".docx"
        document = Document()
        document.add_paragraph(text) 
        out_path = os.path.join(out_dir, filename)
        document.save(out_path)
